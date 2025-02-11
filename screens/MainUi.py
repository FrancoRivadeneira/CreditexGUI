from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.uic import loadUi
import pickle
import socket
import requests
import pandas as pd
import os
import logging
import time
import pygame
from datetime import datetime
import cv2
import threading
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import csv
import numpy as np
pygame.init()
pygame.joystick.init()
from tkinter import Tk, filedialog



import utils.RequestHandler as RequestHandler
from utils.JoystickControlWiFi import JoystickControlWiFi as JoystickControlWiFi
import utils.ControlThread as ControlThread
from utils.VideoThread import VideoThread as VideoThread
from utils.VideoStreamThread import VideoStreamThread as VideoStreamThread
from utils.VideoProcessThread import VideoProcessThread as VideoProcessThread
from utils.VideoThermalThread import VideoThermalThread as VideoThermalThread
from utils.Camera import Camera as Camera

# from utils.VideoStreamThread import VideoStreamThread as VideoStreamThread
exit_flag = threading.Event()

url1_check = "http://192.168.68.102:8080"
url2_check = "http://192.168.68.102:8085"
url3_check = "http://192.168.68.102:8070"

url = "http://192.168.68.102:8080/?action=stream"
url2 = "http://192.168.68.102:8085/?action=stream"
url3 = "http://192.168.68.102:8070/?action=stream"  # Brazo


COORDINATES = []
""" DATOS DE LA CÁMARA """

FRAME_RATE = 10
WIDTH_FRAME_CAM = 640
HEIGHT_FRAME_CAM = 480

USER=''
USERNAME=''
USERCODE=''
USERPICKPATH=''
MAP_NAME = ''
MAP_AREA = ''
MAP_PIC = ''

YAW=0
PITCH=0
ROLL=0
SPEED=0

""" DATOS DEL HARDWARE """
# ROBOT INFO (RASPBERRY REMOTE)
IP_ADDR = "192.168.68.102"  # NUC
IP_ADDR_B = "192.168.0.14"  # "192.168.0.14"
IP_ADDR_CAM = "192.168.23.176"
PORT = 8666

PORT_ANGLE = 5500
PORT_IMU= 9100
PORT_B = 8900
PORT1 = 9692  # Rgb 1
PORT2 = 9992  # Rgb 2
PORT3 = 9500  # Puerto Cámara Brazo
PORT4 = 7002

_FLAG_CONECT = False

FLAG_MAPA= False
FLAG_AUTO= False
is_taking_pictures = False  # determina si empezo el proceso de captura de imagenes
# guarda la direccion mas reciente en la que se grabaron imagenes
picture_directory = None
picture_timer = None  # guarda el temporizador para la captura de imagenes
ini = True  # Bandera que indica cuando se empieze a grabar
is_saving= False
id_gen = 1
id_sec = 2

slider_pos = 0  # guarda la posicion del slider que indica la posicion
counter = 0  # COntador para la pantalla de carga
width = 960  # Ancho del video mostrado en la interfaz
height = 540  # Alto del video mostrado en la interfaz
# Bandera que indica si un set de imagenes ha sido seleccionado como original o como procesado


resources="""QComboBox::drop-down {
    border: 0px;
}
QComboBox::down-arrow {
    image: url(iconos/desplegar.svg);
    height: 40px;
    width: 40px;
    margin-right: 60px;  /* Ajusta el margen según sea necesario */
}
QComboBox {
    padding-right: 50px;  /* Ajusta el padding para que haya espacio para el icono */
    padding-left: 40px;  /* Ajusta el padding para que haya espacio para el icono */
    color:white;
	background-color:#2C2C2F;
	border-radius: 30px;
}
QComboBox QAbstractItemView {
    background: white;  /* Fondo blanco para la lista desplegable */
    color: black;  /* Texto negro para las opciones */
}

"""

class socketReception(QThread):
    _SIGNAL_tramaDatos = pyqtSignal(str)
    def __init__(self,parent):
        super().__init__(None)
        self.falseParent=parent
        
    def run(self):
        self._FLAG_run=True
        clock=pygame.time.Clock()
        while self._FLAG_run and not(exit_flag.is_set()):
            try:
                if not(self.falseParent._FLAG_socketConected):
                    """ Establecemos conexion con el socket para el robot """
                    # Creamos la conexión por socket
                    _VAR_socketRobot= socket.socket(socket.AF_INET,socket.SOCK_STREAM)
                    _VAR_socketRobot.settimeout(10)
                    _VAR_socketRobot.connect((IP_ADDR,PORT))
                    logging.info("Conectado a {} {}!".format(IP_ADDR,PORT))
                    # Modificamos la bandera
                    self.falseParent._FLAG_socketConected=True
                """ Enviamos los datos del joystick """
                _VAR_comando=self.falseParent._VAR_joystickCommand
                _VAR_socketRobot.sendall(pickle.dumps(_VAR_comando))
                """ Recibimos la trama del estado del robot """
                tramaDatos = _VAR_socketRobot.recv(60)
                tramaDatos = pickle.loads(tramaDatos)
                self._SIGNAL_tramaDatos.emit(tramaDatos)
                """ Definimos el tiempo de refresco """
                time.sleep(0.02)
                clock.tick(10)
            except Exception as e:
                logging.info("Se perdió conexión. Volviendo a conectar con el socket [ERROR:{}]".format(e)+"No fuca")
                self.falseParent._FLAG_socketConected=False
                _VAR_socketRobot=None
                time.sleep(0.5)
        """ Cerramos el puerto """
        try:
            _VAR_socketRobot.close()
            logging.info("socket con el robot cerrado!")
        except:
            logging.info("El socket ya se habia cerrado previamente..")
            
    def stop(self):
        self._FLAG_run=False
        self.wait()

class Worker(QtCore.QObject):
    update_label = QtCore.pyqtSignal(str)  # Señal para actualizar el label

    def __init__(self, parent=None):
        super().__init__(parent)

    def imu(self):
        # Aquí va tu código de la función imu
        while True:
            time.sleep(0.05)
            data = f"{ROLL};{PITCH};{YAW};{SPEED}"  # Supongamos que estos son los nuevos datos
            # print(f"llego estoooooaaaaaaaaa: {data}")
            self.update_label.emit(data)  # Emitimos la señal con los nuevos dato



class DimensionUpdater(QObject):
    dimensions_changed = pyqtSignal(int, int)


class MainUi (QtWidgets.QMainWindow):
    """Función Constructor"""

    def __init__(self):
        super(MainUi, self).__init__()

        # Cargamos interfaz principal
        loadUi("C:/Users/Franco Rivadeneira/Desktop/PlantasIndustrialesV2/screens/interfaceMOD.ui", self)
        self.map_df = pd.read_csv('maps.csv')
        """ DEFINICION de FLAG """
        self._FLAG_socketConected = False
        self._FLAG_sshImuConnected = False
        self._FLAG_sshLidarConnected = False
        self._FLAG_recordVideos = False
        """ DEFINICION de Variables """
        # socket
        self._VAR_socketClient = None
        # Video counter
        self._VAR_counterVideoFiles = 0
        # List that contains the VideoWriter to MP4 of the video
        self._VAR_outVideoWrite_list = [None, None]
        self._VAR_tramaDatos = ""
        self._VAR_joystickCommand = ""
        """Se crean los objetos para modificar el tamaño de las cámaras segun se requiera"""
        self.dimension_updater1 = DimensionUpdater()
        self.dimension_updater2 = DimensionUpdater()
        self.dimension_updater3 = DimensionUpdater()
        self.dimension_updater4 = DimensionUpdater()
        self.dimension_updater5 = DimensionUpdater()
        self.dimension_updaterB = DimensionUpdater()
        self.dimension_updater_thermal = DimensionUpdater()
        self.dimension_updaterB_P = DimensionUpdater()
        """Se inicializan los labels e iconos del usuario"""
        self.label_wel.setText(f"Hola {USERNAME}, ¿Que haras hoy?")
        pixmap = QPixmap(USERPICKPATH)
        pixmap = pixmap.scaled(100, 100, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
        icon = QIcon(pixmap)
        self.btn_user.setIcon(icon)
        self.btn_user.setIconSize(QtCore.QSize(250, 250))

        self.thread = QtCore.QThread(self)
        self.worker = Worker()
        self.worker.moveToThread(self.thread)
        self.worker.update_label.connect(self.updateYaw)  # Conectamos la señal a la ranura
        self.thread.started.connect(self.worker.imu)
        self.thread.start()

        #self.setGeometry(0, 0, 1920, 1072)
        # threading.Thread(target=self.dataAng).start()

        # Quitamos el borde y adicionamos fondo transparente
        self.setWindowFlag(QtCore.Qt.FramelessWindowHint)
        self.setAttribute(QtCore.Qt.WA_TranslucentBackground)
       # Se abre maximizando la pantalla
        self.showNormal()
        self.showFullScreen()
        self.initialize_map_frames()
        self.initialize_comboBox()
        # Se verifica que se reconoce el mando, caso contrario el valor de esta variable es 0
        joystick_count = pygame.joystick.get_count()

        # Condicion en caso el joystick este conectado
        if joystick_count != 0:

            # Iniciamos la clase correspondiente al control

            self._THREAD_socketReceptionData = socketReception(self)
            # Check the joystick connection status and update command
            self._THREAD_joytickCommand = JoystickControlWiFi(self)
            self._THREAD_socketReceptionData.start()
            self._THREAD_joytickCommand.start()

            # self.label_18.setStyleSheet(f"color: #006400")
            # self.label_18.setText('Conectado')
    
        # Condicion en caso el Joystick este desconectado
        else:
            """Se define la conexión por WiFi"""
            self._THREAD_socketReceptionData = socketReception(self)
            # Check the joystick connection status and update command
            self._THREAD_joytickCommand = JoystickControlWiFi(self)
            self._THREAD_socketReceptionData.start()
            self._THREAD_joytickCommand.start()

            # Texto que indica la conexion del control
            # self.label_18.setStyleSheet(f"color: #FF0000")
            # self.label_18.setText('Desconectado')
            


        # self.video_thread1 = VideoStreamThread(url,self.dimension_updater1)
        # self.video_thread2 = VideoStreamThread(url2,self.dimension_updater2)
        # self.video_thread3 = VideoStreamThread(url3,self.dimension_updater3)
        # self.video_thread3 = VideoClient(self,2,self.dimension_updater3,IP_ADDR_CAM,PORT3)

        try:
            timeout_seconds = 2
            # Realizar la solicitud y establecer el tiempo de espera
            response = requests.get(url1_check, timeout=timeout_seconds)
            if response.status_code == 200:
                # Si la solicitud fue exitosa, continuar con el proceso de conexión
                self.video_thread = VideoStreamThread(url, self.dimension_updater1)



                # self._THREAD_updateCamera1 = Camera(self, 1, self.dimension_updater)
                # self._THREAD_updateCamera1.change_pixmap_signal.connect(self.dimension_updater1)
            else:
                # Si la solicitud no fue exitosa, imprimir un mensaje de error
                print("Error: No se pudo conectar al URL. Código de estado:", response.status_code)
        except Exception as e:
            logging.exception(e)

        try:
            response = requests.get(url2_check, timeout=timeout_seconds)
            if response.status_code == 200:
                # Si la solicitud fue exitosa, continuar con el proceso de conexión
                self.video_thread2 = VideoStreamThread(url2, self.dimension_updater2)
                # self._THREAD_updateCamera2 = Camera(self, 2, self.dimension_updater)
                # self._THREAD_updateCamera2.change_pixmap_signal.connect(self.dimension_updater2)
                
            else:
                # Si la solicitud no fue exitosa, imprimir un mensaje de error
                print("Error: No se pudo conectar al URL. Código de estado:", response.status_code)

        except Exception as e:
            logging.exception(e)

        try:
            response = requests.get(url3_check, timeout=timeout_seconds)
            if response.status_code == 200:
                # Si la solicitud fue exitosa, continuar con el proceso de conexión
                self.video_thread3 = VideoStreamThread(url3, self.dimension_updater3)
                self.video_threadB = VideoStreamThread(url3, self.dimension_updaterB)
                self.video_threadB_P = VideoProcessThread(self.dimension_updaterB_P)
            else:
                # Si la solicitud no fue exitosa, imprimir un mensaje de error
                print("Error: No se pudo conectar al URL. Código de estado:", response.status_code)
        except Exception as e:
            logging.exception(e)


        # try:
        #     self.video_thread4 = VideoThermalThread(self.dimension_updater4)

        # except ConnectionError as e:
        #     logging.exception(e)
        #     #print(f"No se encontraron cámaras, {e}")



############################################################################################################################################
# BOTONES PARA ABRIR / CERRAR VENTANA

        # Minimize Window
        # self.minimize_window_button.clicked.connect(
        #     lambda: self.showMinimized())

        # Close window
        self.close_window_button.clicked.connect(lambda: self.close())

        # # Close Window Alternativa
        # self.btn_exit.clicked.connect(lambda: self.close())

        # Expand / Restore Window
        # self.max_button.clicked.connect(lambda: self.showFullScreen())

        # # Funcion que aumenta tamano del video con el agrandamiento de pantalla
        # self.max_button.clicked.connect(self.maximize_video)
        

        # Achicar la interfaz
        
        self.btn_toggle.clicked.connect(self.Autonomous)
####################################### CAMBIO DE PESTAÑA #################################################################
        self.btn_general.clicked.connect(self.distribuir_general)
        self.btn_parada.clicked.connect(self.parada)
        self.btn_angulo.clicked.connect(self.distribuir_arm)
        self.btn_document.clicked.connect(self.distribuir_document)
        self.btn_archivos.clicked.connect(self.distribuir_archivos)
        self.btn_posicion.clicked.connect(self.distribuir_posicion)
        self.btn_tumi.clicked.connect(self.distribuir_secret)
        self.btn_reset.clicked.connect(self.reset_button)
        self.btn_estadisticas.clicked.connect(self.GoEstadisticas)
        self.btn_generarReport.clicked.connect(self.GoGenerarReport)
########################################## SLIDERS BRAZO  #################################################################
        # SLIDERS

        # Actualiza Slider
        self.slider_m1.valueChanged.connect(self.update_m1_angle)
        self.slider_m2.valueChanged.connect(self.update_m2_angle)
        self.slider_m3.valueChanged.connect(self.update_m3_angle)
        self.slider_m4.valueChanged.connect(self.update_m4_angle)
        self.slider_m5.valueChanged.connect(self.update_m5_angle)
        self.slider_m6.valueChanged.connect(self.update_m6_angle)
########################################## Galeria  #################################################################
        self.btn_fotos.clicked.connect(lambda: self.load_images_galery("imagenes"))
        self.btn_album.clicked.connect(self.load_files_galery)
        # BOTONES PARA GRABAR
        self.btn_grabar.clicked.connect(self.start_taking_pictures)

        # BOTONES PARA PARAR DE GRABAR
        # self.btn_stop.clicked.connect(self.stop_recording)
        self.btn_test.clicked.connect(self.depuracion)

        # BOTON PARA GENERAR REPORTES

        self.btn_gen_report.clicked.connect(self.generate_report)


        # BOTONES PARA ABRIR ARCHIVOS:
        #self.btn_archivos.clicked.connect(self.open_file_dialog)


        # FUNCION QUE PERMITE ARRASTRAR LA VENTANA POR LA PANTALLA

        def moveWindow(e):
            # Primero verificamos si la ventana esta maximizada
            if self.isFullScreen() == False:  # Not maximized
                # Solo movemos si la ventana no esta agrandada
                # ###############################################
                # Arrastramos la ventana cuando se hace click izquierdo con el mouse
                if e.buttons() == Qt.LeftButton:
                    # Movemos ventana
                    self.move(self.pos() + e.globalPos() - self.clickPosition)
                    self.clickPosition = e.globalPos()
                    e.accept()
        #######################################################################

        # Mover ventana
        self.header_frame.mouseMoveEvent = moveWindow
        self.open_close_side_bar_btn.clicked.connect(lambda: self.slideLeftMenu())
        self.open_close_side_bar_btn.enterEvent = self.open_close_side_bar_btn_enter_event
        self.open_close_side_bar_btn.leaveEvent = self.open_close_side_bar_btn_leave_event
        # INICIAMOS TIMER PARA MOSTRAR SECUENCIA DE IMAGENES

        # Boton Brazo reposo-> Erecto

        self.act1_button.clicked.connect(self.save_arm)
        self.act2_button.clicked.connect(self.atack1)
        self.act3_button.clicked.connect(self.pick1)
        self.act4_button.clicked.connect(self.drop1)
        self.act5_button.clicked.connect(self.stopMotors)
        self.act6_button.clicked.connect(self.leaveMotors)
        self.act7_button.clicked.connect(self.updatePos)
        self.act8_button.clicked.connect(self.send_message)
        self.act10_button.clicked.connect(self.temp)
        self.act12_button.clicked.connect(self.PID)
        self.setPos_button.clicked.connect(self.upick1)

        # BOTONES ACCIONES BRAZO
        # self.act7_button.enterEvent = self.act7_btn_enter_event  # Cambiar Color boton
        # self.act7_button.leaveEvent = self.act7_btn_leave_event

        # Función Click
        self.VisionCamara2.mousePressEvent = lambda event: self.change_camera(event, id_sec)
        self.VisionCamara3.mousePressEvent = lambda event: self.change_camera2(event, id_sec)
        self.VisionCamara4.mousePressEvent = lambda event: self.change_camera3(event, id_sec)
        self.act6_button.setEnabled(True)

    @pyqtSlot(np.ndarray)
    def update_image1(self, cv_img):
        """Updates the image_label with a new opencv image"""
        rgb_image = cv2.cvtColor(cv_img, cv2.COLOR_BGR2RGB)
        height, width, channel = rgb_image.shape
        q_image = QImage(rgb_image.data, width, height, width * channel, QImage.Format_RGB888)        
        scaledImage = q_image.scaled(self.ui.label_camMain.size(), aspectRatioMode=Qt.KeepAspectRatio)    
        pixMap=QPixmap.fromImage(scaledImage)
        self.ui.label_camMain.setPixmap(pixMap)
        self.ui.label_camMain.setAlignment(Qt.AlignCenter)
    
    @pyqtSlot(np.ndarray)
    def update_image2(self, cv_img):
        """Updates the image_label with a new opencv image"""
        rgb_image = cv2.cvtColor(cv_img, cv2.COLOR_BGR2RGB)
        height, width, channel = rgb_image.shape
        q_image = QImage(rgb_image.data, width, height, width * channel, QImage.Format_RGB888)        
        scaledImage = q_image.scaled(self.ui.label_camAruco.size(), aspectRatioMode=Qt.KeepAspectRatio)
        pixMap=QPixmap.fromImage(scaledImage)
        self.ui.label_camAruco.setPixmap(pixMap)
        self.ui.label_camAruco.setAlignment(Qt.AlignCenter)
    
    @QtCore.pyqtSlot(str)
    def updateYaw(self, data):
        
        # Esta es la ranura que actualiza el label
        try:
            roll,pitch,yaw,speed=data.split(';')
        except:
            roll=ROLL
            pitch=PITCH
            yaw=YAW
            speed=SPEED
        # roll,pitch,yaw,speed=data.split(';')
        # roll = int(roll)
        # pitch = int(pitch)
        # yaw = int(yaw)
        speed = int(float(speed))
        # print(type(speed))
        #print(speed)

        # Imprimir los resultados
        #print(f"Roll: {roll}, Pitch: {pitch}, Yaw: {yaw}")
        
        self.label_yaw.setText(f"Yaw: {yaw}")
        self.label_pitch.setText(f"Pitch: {pitch}")
        self.label_roll.setText(f"Roll: {roll}")
        self.speed_bar.setValue(speed)
        self.speed_bar.setFont(QFont('Arial',15,10))
        

        

        # pixmap = QPixmap(f"pitch_images/pitch_nar_30")
        # pixmap = pixmap.scaled(150, 150, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
        # icon = QIcon(pixmap)
        # self.btn_pitch.setIcon(icon)
        # self.btn_pitch.setIconSize(QtCore.QSize(150, 150))

        #self.speed_bar.setAlignment(Qt.AlignCenter)
        self.speed_bar.setFormat(f"{speed}%")
    def sample(self):
        print("aaaaa")
    def parada(self):
        if self.motor_der.isChecked()==True:
            print("Estaba prendido")
            self.motor_der.setChecked(False)
        else:
            print("Estaba apagado")
            self.motor_der.setChecked(True)
        global _FLAG_CONECT

        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = ('192.168.0.10', PORT)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            s.connect(server_address)
            _FLAG_CONECT = True


        try:
            # Enviar datos
            message = '$OAX1Y' #Comando Liberar motores
            print(f"Enviando: {message}")
            s.sendall(message.encode('utf-8'))
            # Buscar respuesta
            data = None
            data = s.recv(256)

            if data:
                #
                # amount_received += len(data)
                data = data.decode('utf-8')
                print(f"Recibido: {data}")
                # motores = data.split(",")
                # if motores[0]=="Apagado":
                #     self.motor_der.setCheck()


                _FLAG_CONECT = False

            else:
                _FLAG_CONECT = False
        except IndexError:
            print("No llegó un dato anguloso")

        except:
            print("Cerrando socket")
            s.close()
            _FLAG_CONECT = False

        """FUNCIONES BRAZO"""
    ############################################### FUNCIONES ComboBoX ######################################################
    def initialize_comboBox(self):
        self.comboMapa.setStyleSheet(resources)
        self.add_items_ComboBox(self.comboMapa,"maps.csv","name")
        self.comboArea.setStyleSheet(resources)
        self.add_items_ComboBox(self.comboArea,"maps.csv","area")
        self.comboTurno.setStyleSheet(resources)
        self.comboCodigo.setStyleSheet(resources)
        self.comboCodigo.addItem(str(USERCODE))
        pixmap = QPixmap(USERPICKPATH)
        pixmap = pixmap.scaled(360, 360, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
        icon = QIcon(pixmap)
        self.btn_profile_report.setIcon(icon)
        self.btn_profile_report.setIconSize(QtCore.QSize(400, 250))
    ############################################### FUNCIONES MAPA  #########################################################
    def initialize_map_frames(self):
        for index, row in self.map_df.iterrows():
            name = row['name']
            image_path = row['photo']
            area = row['area']
            self.create_map_frame(name,area,image_path)
    def create_map_frame(self, name, area ,image_path):
        frame = QFrame()
        frame.setMinimumSize(260, 350)
        frame.setMaximumSize(260, 350)

        layout = QVBoxLayout()
        button = QPushButton()
        button.setMinimumSize(250, 250)
        button.setMaximumSize(250, 250)
        pixmap = QPixmap(image_path)
        pixmap = pixmap.scaled(250, 250, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
        icon = QIcon(pixmap)
        button.setIcon(icon)
        button.setIconSize(QtCore.QSize(250, 250))

        label_name = QLabel(str(name))  # Convertir a str si es necesario
        label_area = QLabel(str(area))  # Convertir a str si es necesario
        
        font = QFont()
        font.setFamily('MS Shell Dlg 2')
        font.setPointSize(14)
        
        label_name.setFont(font)
        label_area.setFont(font)
        
        label_name.setStyleSheet("color: white; text-align: center;")
        label_area.setStyleSheet("color: white; text-align: center;")
        label_name.setAlignment(QtCore.Qt.AlignCenter)
        label_area.setAlignment(QtCore.Qt.AlignCenter)
        
        # Set maximum width for labels
        label_name.setWordWrap(True)
        label_area.setWordWrap(True)

        layout.addWidget(button)
        layout.addWidget(label_name)
        layout.addWidget(label_area)
        frame.setLayout(layout)
        button.clicked.connect(lambda checked, user=name, area_map=area,img=image_path: self.GoGeneral(user,area_map, img))
        # Insert frame to the left of frame_mapgen
        layout_usuarios = self.frame_mapas_list.layout()
        index_create = layout_usuarios.indexOf(self.frame_mapgen)
        layout_usuarios.insertWidget(index_create, frame)
    ############################################### FUNCIONES GALERY ########################################################

        
    def load_images_galery(self,folder_path="imagenes",isAlbum=False):
        if not isAlbum:
            self.btn_fotos.setStyleSheet("QPushButton { color: black; background-color: #EA6C36; border-radius:20px; }")
            self.btn_album.setStyleSheet("QPushButton { color: white; background-color: #0F0F0F; border-radius:20px; }")
        else:
            self.btn_album.setStyleSheet("QPushButton { color: black; background-color: #EA6C36; border-radius:20px; }")
            self.btn_fotos.setStyleSheet("QPushButton { color: white; background-color: #0F0F0F; border-radius:20px; }")
        # Check if the folder_path exists
        if not os.path.exists(folder_path):
            print(f"The folder {folder_path} does not exist.")
            return

        # Create a scroll area
        scroll_area = QScrollArea(self)
        scroll_area.setWidgetResizable(True)

        # Create a container widget and set it as the widget for the scroll area
        container_widget = QWidget()
        scroll_area.setWidget(container_widget)

        # Create a grid layout for the container widget
        layout = QGridLayout(container_widget)
        container_widget.setLayout(layout)

        # Recursive function to find all image files in the folder and its subfolders
        def find_image_files(folder):
            image_files = []
            for root, dirs, files in os.walk(folder):
                for file in files:
                    if file.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                        image_files.append(os.path.join(root, file))
            return image_files

        # Find all image files in the folder and its subfolders
        image_files = find_image_files(folder_path)

        row, col = 0, 0
        for i, file_path in enumerate(image_files):
            print(f"Processing file: {file_path}")
            
            pixmap = QPixmap(file_path)
            if pixmap.isNull():
                print(f"Failed to load image: {file_path}")
                continue

            if col == 4:  # Move to the next row after 4 images
                col = 0
                row += 1

            frame = QFrame()
            frame.setMinimumSize(260, 260)
            frame.setMaximumSize(260, 260)
            
            layout_frame = QVBoxLayout()
            button = QPushButton()
            button.setMinimumSize(250, 250)
            button.setMaximumSize(250, 250)
            pixmap = pixmap.scaled(250, 250, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
            icon = QIcon(pixmap)
            button.setIcon(icon)
            button.setIconSize(QtCore.QSize(250, 250))
            
            layout_frame.addWidget(button)
            frame.setLayout(layout_frame)
            
            layout.addWidget(frame, row, col, 1, 1)
            col += 1

        # Ensure frame_galery is properly configured
        if not hasattr(self, 'frame_galery') or not self.frame_galery:
            self.frame_galery = QFrame(self)
            main_layout = QVBoxLayout(self.frame_galery)
            self.frame_galery.setLayout(main_layout)
        else:
            main_layout = self.frame_galery.layout()

        # Clear existing layout from frame_galery
        while main_layout.count():
            item = main_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()

        # Add the scroll area to the frame_galery
        main_layout.addWidget(scroll_area)

    def load_files_galery(self):
        self.btn_album.setStyleSheet("QPushButton { color: black; background-color: #EA6C36; border-radius:20px; }")
        self.btn_fotos.setStyleSheet("QPushButton { color: white; background-color: #0F0F0F; border-radius:20px; }")
        folder_path="imagenes"
        # Check if the folder_path exists
        if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
            print(f"The folder {folder_path} does not exist or is not a directory.")
            return

        # Create a scroll area
        scroll_area = QScrollArea(self)
        scroll_area.setWidgetResizable(True)

        # Create a container widget and set it as the widget for the scroll area
        container_widget = QWidget()
        scroll_area.setWidget(container_widget)

        # Create a grid layout for the container widget
        layout = QGridLayout(container_widget)
        container_widget.setLayout(layout)

        # List all subdirectories in the folder
        subdirectories = [name for name in os.listdir(folder_path) if os.path.isdir(os.path.join(folder_path, name))]

        row, col = 0, 0
        for i, subdir in enumerate(subdirectories):
            print(f"Processing folder: {subdir}")
            
            # Create button
            button = QPushButton()
            button.setIcon(QIcon('iconos/carpeta.svg'))  # Replace 'carpeta.svg' with your icon file
            button.setIconSize(QtCore.QSize(300, 300))  # Adjust icon size as needed
            button.clicked.connect(lambda checked, img=f"imagenes/{subdir}": self.load_images_galery(img,True))
            
            
            # Create label with folder name
            label = QLabel(subdir)
            label.setAlignment(QtCore.Qt.AlignCenter)
            label.setStyleSheet("QLabel { color: white; font-size: 30px; }")

            # Create frame to hold button and label
            frame = QFrame()
            frame.setMaximumHeight(350)
            frame_layout = QVBoxLayout()
            frame_layout.addWidget(button)
            frame_layout.addWidget(label)
            frame.setLayout(frame_layout)
            
            # Add frame to the layout
            layout.addWidget(frame, row, col, 1, 1)

            col += 1
            if col == 4:  # Move to the next row after 4 columns
                col = 0
                row += 1

        # Ensure frame_galery is properly configured
        if not hasattr(self, 'frame_galery') or not self.frame_galery:
            self.frame_galery = QFrame(self)
            main_layout = QVBoxLayout(self.frame_galery)
            self.frame_galery.setLayout(main_layout)
        else:
            main_layout = self.frame_galery.layout()

        # Clear existing layout from frame_galery
        while main_layout.count():
            item = main_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()

        # Add the scroll area to the frame_galery
        main_layout.addWidget(scroll_area)
    ############################################## FUNCION MENU DESPLEGABLE  ###############################################

    def open_close_side_bar_btn_enter_event(self, event):

        width = self.slide_menu_container.width()
        if width <= 0:
            # Change the button color when hovered
            self.open_close_side_bar_btn.setIcon(
                QtGui.QIcon(u"Iconos/menu2.png"))

        else:
            # Change the button color when hovered
            self.open_close_side_bar_btn.setIcon(
                QtGui.QIcon(u"Iconos/left2.png"))
    def open_close_side_bar_btn_leave_event(self, event):

        width = self.slide_menu_container.width()
        if width <= 0:
            # Change the button color when hovered
            self.open_close_side_bar_btn.setIcon(
                QtGui.QIcon(u"Iconos/menu.png"))
        else:
            self.open_close_side_bar_btn.setIcon(
                QtGui.QIcon(u"Iconos/left.png"))
    ############################################## FUNCIONES BRAZO ###############################################
    def upick1(self):
        global _FLAG_CONECT

        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            _VAR_socketB = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = (IP_ADDR_B, PORT_B)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            _VAR_socketB.connect(server_address)
            _FLAG_CONECT = True

        try:
            # Enviar datos
            message = f'$OAX2U1'
            print(f"Enviando: {message}")
            _VAR_socketB.sendall(message.encode('utf-8'))
            _FLAG_CONECT = False

        except:
            print("Cerrando socket")
            _VAR_socketB.close()
            _FLAG_CONECT = False
    def send_message(self):
            global _FLAG_CONECT

            if not _FLAG_CONECT:

                # Crear un socket TCP/IP
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                # Conectar el socket al puerto donde el servidor está escuchando
                server_address = (IP_ADDR_B, PORT_B)
                print(f"Conectando a {server_address[0]} puerto {server_address[1]}")
                s.connect(server_address)
                _FLAG_CONECT = True

            try:
                # Enviar datos
                message = '$OAX2R0'
                print(f"Enviando: {message}")
                s.sendall(message.encode('utf-8'))
                # Buscar respuesta
                data = None
                data = s.recv(256)

                if data:
                    #
                    data = data.decode('utf-8')
                    print(f"Recibido: {data}")
                    angulos = data.split(",")
                    angulo_0 = int(angulos[0])/100
                    self.rep_ang1.setText(f"Angle 1: {angulo_0}°")
                    angulo_1 = int(angulos[1])/100
                    self.rep_ang2.setText(f"Angle 2: {angulo_1}°")
                    angulo_2 = int(angulos[2])/100
                    self.rep_ang3.setText(f"Angle 3: {angulo_2}°")
                    angulo_3 = int(angulos[3])/100
                    self.rep_ang4.setText(f"Angle 4: {angulo_3}°")
                    angulo_4 = int(angulos[4])/100
                    self.rep_ang5.setText(f"Angle 5: {angulo_4}°")
                    angulo_5 = int(angulos[5])/100
                    self.rep_ang6.setText(f"Angle 6: {angulo_5}°")


                    
                    self.angle_m1.setText(f"{round(angulo_0,0)}"[:-2]+"°")
                    self.angle_m2.setText(f"{round(angulo_1,0)}"[:-2]+"°")
                    self.angle_m3.setText(f"{round(angulo_2,0)}"[:-2]+"°")
                    self.angle_m4.setText(f"{round(angulo_3,0)}"[:-2]+"°")
                    self.angle_m5.setText(f"{round(angulo_4,0)}"[:-2]+"°")
                    self.angle_m6.setText(f"{round(angulo_5,0)}"[:-2]+"°")

                    """Extracción de °"""
                    angle1_completo = self.angle_m1.text()
                    angle1 = angle1_completo[:-1]  # Se le quita el simbolo "°"
                    angle2_completo = self.angle_m2.text()
                    angle2 = angle2_completo[:-1]  # Se le quita el simbolo "°"
                    angle3_completo = self.angle_m3.text()
                    angle3 = angle3_completo[:-1]  # Se le quita el simbolo "°"
                    angle4_completo = self.angle_m4.text()
                    angle4 = angle4_completo[:-1]  # Se le quita el simbolo "°"
                    angle5_completo = self.angle_m5.text()
                    angle5 = angle5_completo[:-1]  # Se le quita el simbolo "°"
                    angle6_completo = self.angle_m6.text()
                    angle6 = angle6_completo[:-1]  # Se le quita el simbolo "°"

                    """Casteo de angulos"""

                    
                    angle1 = int(angle1)
                    angle2 = int(angle2)
                    angle3 = int(angle3)
                    angle4 = int(angle4)
                    angle5 = int(angle5)
                    angle6 = int(angle6)

                    print(f"el numero es:{angle1}+{angle2}+{angle3}+{angle4}+{angle5}+{angle6}")

                    self.slider_m1.setValue(angle1)
                    self.slider_m2.setValue(angle2)
                    self.slider_m3.setValue(angle3)
                    self.slider_m4.setValue(angle4)
                    self.slider_m5.setValue(angle5)
                    self.slider_m6.setValue(angle6)


                    _FLAG_CONECT = False

                else:
                    _FLAG_CONECT = False
            except IndexError:
                print("No llegó un dato anguloso")

            except:
                print("Cerrando socket")
                s.close()
                _FLAG_CONECT = False
    def PID(self):
            global _FLAG_CONECT

            if not _FLAG_CONECT:

                # Crear un socket TCP/IP
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                # Conectar el socket al puerto donde el servidor está escuchando
                server_address = (IP_ADDR_B, PORT_B)
                print(f"Conectando a {server_address[0]} puerto {server_address[1]}")
                s.connect(server_address)
                _FLAG_CONECT = True

            try:
                # Enviar datos
                message = '$OAX9W1S030030\n$OAX9W1P150150'
                print(f"Enviando: {message}")
                s.sendall(message.encode('utf-8'))
                # Buscar respuesta
                data = None
                data = s.recv(256)

                if data:
                    #
                    data = data.decode('utf-8')
                    print(f"Recibido: {data}")
                    _FLAG_CONECT = False

                else:
                    print()
                    _FLAG_CONECT = False
            except IndexError:
                print("No llegó un dato anguloso")

            except:
                print("Cerrando socket")
                s.close()
                _FLAG_CONECT = False
    def temp(self):
            global _FLAG_CONECT

            if not _FLAG_CONECT:

                # Crear un socket TCP/IP
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                # Conectar el socket al puerto donde el servidor está escuchando
                server_address = (IP_ADDR_B, PORT_B)
                print(f"Conectando a {server_address[0]} puerto {server_address[1]}")
                s.connect(server_address)
                _FLAG_CONECT = True

            try:
                # Enviar datos
                message = '$OAX2P1'
                print(f"Enviando: {message}")
                s.sendall(message.encode('utf-8'))
                # Buscar respuesta
                data = None
                data = s.recv(256)

                if data:
                    #
                    data = data.decode('utf-8')
                    print(f"Recibido: {data}")
                    _FLAG_CONECT = False

                else:
                    print()
                    _FLAG_CONECT = False
            except IndexError:
                print("No llegó un dato anguloso")

            except:
                print("Cerrando socket")
                s.close()
                _FLAG_CONECT = False
    def leaveMotors(self):
        global _FLAG_CONECT

        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            _VAR_socketB = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = (IP_ADDR_B, PORT_B)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            _VAR_socketB.connect(server_address)
            _FLAG_CONECT = True

        try:
            # Enviar datos
            message = f'$OAX2S1\n$OAX2S2\n$OAX2S3\n$OAX2S4\n$OAX2S5\n$OAX2S6\n'
            print(f"Enviando: {message}")
            _VAR_socketB.sendall(message.encode('utf-8'))
            _FLAG_CONECT = False

        except:
            print("Cerrando socket")
            _VAR_socketB.close()
            _FLAG_CONECT = qFormatLogMessage
    def stopMotors(self):  # Detiene todos los motores
        global _FLAG_CONECT

        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            _VAR_socketB = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = (IP_ADDR_B, PORT_B)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            _VAR_socketB.connect(server_address)
            _FLAG_CONECT = True

        try:
            # Enviar datos
            message = f'$OAX5E1009\n$OAX5E2009\n$OAX5E3009\n$OAX5E4009\n$OAX5E5009\n$OAX5E6009\n'
            print(f"Enviando: {message}")
            _VAR_socketB.sendall(message.encode('utf-8'))
            _FLAG_CONECT = False

        except:
            print("Cerrando socket")
            _VAR_socketB.close()
            _FLAG_CONECT = False
    def updatePos(self):
        time.sleep(0.1)
        if not self.act4_button.isEnabled():  # Verificar si el botón está habilitado
            return
        self.act4_button.setEnabled(False)
        # puerto_serie = serial.Serial('COM8',115200)

        """Extracción de °"""
        angle1_completo = self.angle_m1.text()
        angle1 = angle1_completo[:-1]  # Se le quita el simbolo "°"
        angle2_completo = self.angle_m2.text()
        angle2 = angle2_completo[:-1]  # Se le quita el simbolo "°"
        angle3_completo = self.angle_m3.text()
        angle3 = angle3_completo[:-1]  # Se le quita el simbolo "°"
        angle4_completo = self.angle_m4.text()
        angle4 = angle4_completo[:-1]  # Se le quita el simbolo "°"
        angle5_completo = self.angle_m5.text()
        angle5 = angle5_completo[:-1]  # Se le quita el simbolo "°"
        angle6_completo = self.angle_m6.text()
        angle6 = angle6_completo[:-1]  # Se le quita el simbolo "°"

        """Casteo de angulos"""
        angle1 = int(angle1)  # Convertir a entero si es positivo
        angle2 = int(angle2)
        angle3 = int(angle3)
        angle4 = int(angle4)
        angle5 = int(angle5)
        print(angle5)
        angle6 = int(angle6)



        angle1 = f'{"-" if angle1 < 0 else "+"}{abs(angle1):03}'
        angle2 = f'{"-" if angle2 < 0 else "+"}{abs(angle2):03}'
        angle3 = f'{"-" if angle3 < 0 else "+"}{abs(angle3):03}'
        angle4 = f'{"-" if angle4 < 0 else "+"}{abs(angle4):03}'
        angle5 = f'{"-" if angle5 < 0 else "+"}{abs(angle5):03}'
        angle6 = f'{"-" if angle6 < 0 else "+"}{abs(angle6):03}'

        print(angle5)
        self.act4_button.setEnabled(True)
        clock = pygame.time.Clock
        global _FLAG_CONECT

        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            _VAR_socketB = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = (IP_ADDR_B, PORT_B)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            _VAR_socketB.connect(server_address)
            _FLAG_CONECT = True

        try:
            # Enviar datos
            message = f'$OAX7Q1{angle1}1\n$OAX7Q2{angle2:03}1\n$OAX7Q3{angle3:03}1\n$OAX7Q4{angle4:03}1\n$OAX7Q5{angle5:03}1\n$OAX7Q6{angle6:03}1\n'
            print(f"Enviando: {message}")
            _VAR_socketB.sendall(message.encode('utf-8'))
            _FLAG_CONECT = False

        except:
            print("Cerrando socket")
            _VAR_socketB.close()
            _FLAG_CONECT = False

        self._FLAG_run = True

        """ Cerramos el puerto """
        try:
            _VAR_socketB.close()
            logging.info("socket con el robot cerrado!")
        except:
            logging.info("El socket ya se habia cerrado previamente..")
    def drop1(self):
        global _FLAG_CONECT

        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            _VAR_socketB = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = (IP_ADDR_B, PORT_B)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            _VAR_socketB.connect(server_address)
            _FLAG_CONECT = True

        try:
            # Enviar datos
            message = f'$OAX2J2'#$OAX2J2
            print(f"Enviando: {message}")
            _VAR_socketB.sendall(message.encode('utf-8'))
            _FLAG_CONECT = False

        except:
            print("Cerrando socket")
            _VAR_socketB.close()
            _FLAG_CONECT = False
    def pick1(self):
        global _FLAG_CONECT

        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            _VAR_socketB = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = (IP_ADDR_B, PORT_B)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            _VAR_socketB.connect(server_address)
            _FLAG_CONECT = True

        try:
            # Enviar datos
            message = f'$OAX2J1'
            print(f"Enviando: {message}")
            _VAR_socketB.sendall(message.encode('utf-8'))
            _FLAG_CONECT = False

        except:
            print("Cerrando socket")
            _VAR_socketB.close()
            _FLAG_CONECT = False
    def atack1(self):
        global _FLAG_CONECT
        self.act3_button.setEnabled(True)
        self.act4_button.setEnabled(True)

        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            _VAR_socketB = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = (IP_ADDR_B, PORT_B)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            _VAR_socketB.connect(server_address)
            _FLAG_CONECT = True

        try:
            # Enviar datos
            message = f'$OAX2JA'
            print(f"Enviando: {message}")
            _VAR_socketB.sendall(message.encode('utf-8'))
            _FLAG_CONECT = False

        except:
            print("Cerrando socket")
            _VAR_socketB.close()
            _FLAG_CONECT = False

    def save_arm(self):
        global _FLAG_CONECT

        # self.act3_button.setEnabled(False)
        # self.act4_button.setEnabled(False)

        self.angle_m1.setText("0°")
        self.angle_m2.setText("0°")
        self.angle_m3.setText("0°")
        self.angle_m4.setText("0°")
        self.angle_m5.setText("0°")
        self.angle_m6.setText("0°")
        self.slider_m1.setValue(0)
        self.slider_m2.setValue(0)
        self.slider_m3.setValue(0)
        self.slider_m4.setValue(0)
        self.slider_m5.setValue(0)
        self.slider_m6.setValue(0)


        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            _VAR_socketB = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = (IP_ADDR_B, PORT_B)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            _VAR_socketB.connect(server_address)
            _FLAG_CONECT = True

        try:
            # Enviar datos
            message = f'$OAX2JG'
            print(f"Enviando: {message}")
            _VAR_socketB.sendall(message.encode('utf-8'))
            _FLAG_CONECT = False

        except:
            print("Cerrando socket")
            _VAR_socketB.close()
            _FLAG_CONECT = False

    def update_m1_angle(self, event):
        self.angle_m1.setText(str(event) + '°')

    def update_m2_angle(self, event):
        self.angle_m2.setText(str(event) + '°')

    def update_m3_angle(self, event):
        self.angle_m3.setText(str(event) + '°')

    def update_m4_angle(self, event):
        self.angle_m4.setText(str(event) + '°')

    def update_m5_angle(self, event):
        self.angle_m5.setText(str(event) + '°')

    def update_m6_angle(self, event):
        self.angle_m6.setText(str(event) + '°')

    ############################################## FUNCIONES CAMARA  ###############################################
    def reset_cam_position(self):
        global id_sec
        global id_gen
        self.video_thread.update_frame.disconnect()
        self.video_thread2.update_frame.disconnect()
        self.video_thread3.update_frame.disconnect()
        self.video_thread4.update_frame.disconnect()

        # NO TOCAR
        self.dimension_updater1.dimensions_changed.emit(600, 450)
        self.video_thread.update_frame.connect(self.display_resized_frame1)
        self.video_thread.start()
        self.dimension_updater2.dimensions_changed.emit(200, 150)
        self.video_thread2.update_frame.connect(self.display_resized_frame2)
        self.video_thread2.start()
        self.dimension_updater3.dimensions_changed.emit(200, 150)
        self.video_thread3.update_frame.connect(self.display_resized_frame3)
        self.video_thread3.start()
        self.dimension_updater4.dimensions_changed.emit(200, 150)
        self.video_thread4.update_frame.connect(self.display_resized_frame4)
        self.video_thread4.start()

        self.label_1st.setText('Frontal Camera')
        self.label_2nd.setText('Rear Camera')
        self.label_3rd.setText('Arm Camera')
        self.label_4th.setText('Thermal Camera')

        id_sec = 1
        id_gen = 2

    def change_camera(self, event, id):
        global id_sec
        global id_gen
        global width
        global height
        print(id_gen)
        print(id)

        if picture_directory:
            print("funcionaaaaaaaaaaa"+str(id_gen))

            self.video_thread.update_frame.disconnect()
            self.video_thread2.update_frame.disconnect()

            if id_gen == 1:
                if id_sec == 2:
                    self.dimension_updater2.dimensions_changed.emit(600, 450)
                    self.video_thread2.update_frame.connect(
                        self.display_resized_frame1)
                    self.video_thread2.start()
                    print("llegue")

                    # NO TOCAR
                    self.dimension_updater1.dimensions_changed.emit(200, 150)
                    self.video_thread.update_frame.connect(
                        self.display_resized_frame2)
                    self.label_1st.setText('Rear Camera')
                    self.label_2nd.setText('Frontal Camera')
                    self.video_thread.start()
                    print("LLEGUE")

                    id_sec = 1
                    id_gen = 2
                    print("id_general= "+str(id_gen))
                    print("id_secundario= "+str(id_sec))
                # La camara 1 es la vision general
            elif id_gen == 2:
                if id_sec == 1:

                    self.dimension_updater2.dimensions_changed.emit(200, 150)
                    self.video_thread2.update_frame.connect(
                        self.display_resized_frame2)
                    self.label_2nd.setText('Rear Camera')

                    self.label_1st.setText('Frontal Camera')
                    self.video_thread2.start()

                    self.dimension_updater1.dimensions_changed.emit(600, 450)
                    self.video_thread.update_frame.connect(
                        self.display_resized_frame1)
                    self.video_thread.start()

                    # Establece la altura deseada para la cámara 1
                    id_sec = 2
                    id_gen = 1
                    print("id_general= "+str(id_gen))
                    print("id_secundario= "+str(id_sec))
                # La camara 2 es la vision general
            # Connect the signal to update the label with the video frame
            # Start the picture thread

    def change_camera2(self, event, id):
        global id_sec
        global id_gen
        global width
        global height
        print(id_gen)
        print(id)

        id_sec = 3

        if picture_directory:
            self.reset_cam_position()


            if id_gen == 1:
                if id_sec == 3:
                    self.dimension_updater3.dimensions_changed.emit(600, 450)
                    self.video_thread3.update_frame.connect(self.display_resized_frame1)
                    self.video_thread3.start()
                    print("llegue")

                    # NO TOCAR
                    self.dimension_updater1.dimensions_changed.emit(200, 150)
                    self.video_thread.update_frame.connect(self.display_resized_frame3)
                    self.video_thread.start()
                    print("LLEGUE")

                    self.label_1st.setText('Arm Camera')
                    self.label_3rd.setText('Frontal Camera')

                    id_sec = 1
                    id_gen = 3
                    print("id_general= "+str(id_gen))
                    print("id_secundario= "+str(id_sec))
                # La camara 1 es la vision general
            elif id_gen == 3:
                if id_sec == 1:

                    self.dimension_updater3.dimensions_changed.emit(200, 150)
                    self.video_thread3.update_frame.connect(self.display_resized_frame3)
                    self.video_thread3.start()

                    self.dimension_updater1.dimensions_changed.emit(600, 450)
                    self.video_thread.update_frame.connect(self.display_resized_frame1)
                    self.video_thread.start()

                    self.label_3rd.setText('Arm Camera')
                    self.label_1st.setText('Frontal Camera')
                    # Establece la altura deseada para la cámara 1
                    id_sec = 3
                    id_gen = 1
                    print("id_general= "+str(id_gen))
                    print("id_secundario= "+str(id_sec))
                # La camara 2 es la vision general
            # Connect the signal to update the label with the video frame
            # Start the picture thread

    def change_camera3(self, event, id):
        global id_sec
        global id_gen
        global width
        global height
        print(id_gen)
        print(id)

        id_sec = 4

        if picture_directory:
            self.reset_cam_position()


            if id_gen == 1:
                if id_sec == 4:
                    self.dimension_updater4.dimensions_changed.emit(600, 450)
                    self.video_thread4.update_frame.connect(self.display_resized_frame1)
                    self.video_thread4.start()
                    print("llegue")

                    # NO TOCAR
                    self.dimension_updater1.dimensions_changed.emit(200, 150)
                    self.video_thread.update_frame.connect(self.display_resized_frame4)
                    self.video_thread.start()
                    print("LLEGUE")

                    self.label_1st.setText('Thermal Camera')
                    self.label_4th.setText('Frontal Camera')

                    id_sec = 1
                    id_gen = 4
                    print("id_general= "+str(id_gen))
                    print("id_secundario= "+str(id_sec))
                # La camara 1 es la vision general
            elif id_gen == 4:
                if id_sec == 1:

                    self.dimension_updater4.dimensions_changed.emit(200, 150)
                    self.video_thread4.update_frame.connect(self.display_resized_frame4)
                    self.video_thread4.start()

                    self.dimension_updater1.dimensions_changed.emit(600, 450)
                    self.video_thread.update_frame.connect(self.display_resized_frame1)
                    self.video_thread.start()

                    self.label_4th.setText('Thermal Camera')
                    self.label_1st.setText('Frontal Camera')
                    # Establece la altura deseada para la cámara 1
                    id_sec = 4
                    id_gen = 1
                    print("id_general= "+str(id_gen))
                    print("id_secundario= "+str(id_sec))
                # La camara 2 es la vision general
            # Connect the signal to update the label with the video frame
            # Start the picture thread


    def update_label(self, data):
        self.label_21.setText(data)

#########################################################################################################################################################################

# Funciones para abrir y procesar archivos


    def open_file_dialog(self):
        # Open a file dialog window

        global fn
        global nega_fn
        global current_index
        current_index = 0

        # Activamos botones
        # self.btn_procesado.setEnabled(False)
        # self.btn_original.setEnabled(False)
        # self.btn_procesar.setEnabled(False)
        # self.btn_play.setEnabled(False)
        # self.btn_pause.setEnabled(False)

        # Pedimos a usuario que seleccione el archivo deseaado
        filename, _ = QFileDialog.getOpenFileName(
            self, 'Open Image', '', 'Image files (*.png *.jpg *.jpeg)')

        file_extension = os.path.splitext(filename)[1]

        # Removemos el punto de la extension del archivo
        file_extension = file_extension[1:]

        # Solo confirmmamos si se ha seleccionado una imagen
        if (file_extension == 'jpg') or (file_extension == 'png'):
            pass

            # self.btn_original.setEnabled(True)
            # self.btn_procesar.setEnabled(True)
            # self.btn_play.setEnabled(True)
            # self.btn_pause.setEnabled(True)
            # self.btn_next.setEnabled(True)
            # self.btn_prev.setEnabled(True)

        # Guardamos ruta de archivo en la variable local
        fn = filename

        # A partir de la imagen seleccionada obtenemos el directorio de las imagenes.
        self.original_folder = os.path.dirname(fn)
        self.original_list = sorted(os.listdir(self.original_folder))

        # En caso se seleccione imagen procesada, se verifica si esta carpeta esta debidamente guardada
        nega_folder = os.path.dirname((os.path.dirname(fn))) + '/Nega'

        # Si existe y esa debidamente guardada, la mostramos.
        if os.path.exists(nega_folder):
            # Crea la carpeta de imagenes negativas

            file_extension = fn.replace('.jpg', '')
            nega_fn = nega_folder + '/' + \
                os.path.basename(file_extension) + '_nega.jpg'
            self.nega_folder = os.path.dirname(nega_fn)
            self.nega_list = sorted(os.listdir(self.nega_folder))

            if file_extension == 'jpg':
                pass

                # self.btn_procesar.setEnabled(False)
                # self.btn_procesado.setEnabled(True)
                # self.btn_next.setEnabled(True)
                # self.btn_prev.setEnabled(True)

        print(fn)

        # Entonces mostramos la imagen designada en la label correspondiente
        pixmap = QtGui.QPixmap(filename)
        self.label_24.setPixmap(pixmap)
        self.label_24.setScaledContents(True)

        # De igual manera mostramos la posicion y la cantidad de fallas
        # Obtenemos el nomre de la imagen donde tendremos la informacion
        name_1 = os.path.basename(filename)
        name = os.path.splitext(name_1)[0]  # Quitamos la extension
        words = name.split('_')

        # Seleccionamos el folder en el que se encuentra guardada la imagen
        folder_name = os.path.basename(os.path.dirname(filename))

        # Indicamos en un texto el folder al que pertenece
        self.label_23.setText(folder_name)

    def display_frame(self, frame):
        self.VisionCamara5.setPixmap(frame)

    def start_taking_pictures(self):
        global is_taking_pictures
        global picture_directory
        global height
        global width
        global ini

        ini = False

        # Seleccionamos la carpeta donde se guardarán las imágenes.
        picture_directory = QtWidgets.QFileDialog.getExistingDirectory(
            None, 'Select Picture Directory')

        # Creamos la carpeta Original si es que no existe.
        if not os.path.exists(picture_directory + '/Original'):
            os.makedirs(picture_directory + '/Original')

        folder_original = picture_directory + '/Original'

        if picture_directory:
            try:
                timeout_seconds = 1
                # Realizar la solicitud y establecer el tiempo de espera
                response = requests.get(url1_check, timeout=timeout_seconds)
                if response.status_code == 200:
                    # Si la solicitud fue exitosa, continuar con el proceso de conexión
                    self.video_thread = VideoStreamThread(url, self.dimension_updater1)
                    self.dimension_updater1.dimensions_changed.emit(600, 450)
                    self.video_thread.update_frame.connect(self.display_resized_frame1)
                    self.video_thread.start()

                else:
                    # Si la solicitud no fue exitosa, imprimir un mensaje de error
                    print("Error: No se pudo conectar al URL. Código de estado:", response.status_code)
            except requests.exceptions.Timeout:
                # Manejar el caso en el que la conexión tarda demasiado en establecerse
                print(f"Error: Tiempo de espera agotado para la conexión a {url}")
            except requests.exceptions.RequestException as e:
                # Manejar excepciones relacionadas con solicitudes de red, como conexión fallida o tiempo de espera
                print(f"Error al realizar la solicitud a {url}: {e}")
            except Exception as e:
                logging.exception(e)

            try:
                response = requests.get(url2_check, timeout=timeout_seconds)
                if response.status_code == 200:
                    # Si la solicitud fue exitosa, continuar con el proceso de conexión


                    self.video_thread2 = VideoStreamThread(url2, self.dimension_updater2)
                    self.dimension_updater2.dimensions_changed.emit(200, 150)
                    self.video_thread2.update_frame.connect(self.display_resized_frame2)
                    self.video_thread2.start()
                else:
                    # Si la solicitud no fue exitosa, imprimir un mensaje de error
                    print("Error: No se pudo conectar al URL. Código de estado:", response.status_code)

            except Exception as e:
                logging.exception(e)

            try:
                response = requests.get(url3_check, timeout=timeout_seconds)
                if response.status_code == 200:
                    # Si la solicitud fue exitosa, continuar con el proceso de conexión
                    self.video_thread3 = VideoStreamThread(url3, self.dimension_updater3)
                    self.dimension_updater3.dimensions_changed.emit(200, 150)
                    self.video_thread3.update_frame.connect(self.display_resized_frame3)
                    self.video_thread3.start()
                    self.video_threadB = VideoStreamThread(url3, self.dimension_updaterB)
                    self.dimension_updaterB.dimensions_changed.emit(400, 300)
                    self.video_threadB.update_frame.connect(self.display_resized_frameB)
                    self.video_threadB.start()
                    self.video_threadB_P = VideoProcessThread(self.dimension_updaterB_P)
                    self.dimension_updaterB_P.dimensions_changed.emit(400, 300)
                    self.video_threadB_P.update_frame.connect(self.display_resized_frameB_P)
                    self.video_threadB_P.start()
                else:
                    # Si la solicitud no fue exitosa, imprimir un mensaje de error
                    print("Error: No se pudo conectar al URL. Código de estado:", response.status_code)
            except Exception as e:
                logging.exception(e)


            try:
                # Utiliza una instancia de dimension_updater para video_thread4
                self.video_thread4 = VideoThermalThread(self.dimension_updater4)
                self.dimension_updater4.dimensions_changed.emit(200, 150)
                self.video_thread4.update_frame.connect(self.display_resized_frame4)
                self.video_thread4.start()

                # # Utiliza una instancia de dimension_updater_thermal para video_thread_thermal
                # self.video_thread_thermal = VideoThermalThread(self.dimension_updater_thermal)
                # self.dimension_updater_thermal.dimensions_changed.emit(400, 300)
                # self.video_thread_thermal.update_frame.connect(self.display_resized_frame_thermal)
                # self.video_thread_thermal.start()

            except ConnectionError as e:
                logging.exception(e)
                print(f"No se encontraron cámaras, {e}")


            # Configura las dimensiones para la cámara 2 y luego inicia el hilo
            is_taking_pictures = True

    def display_resized_frame1(self, frame):
        self.VisionCamara1.setPixmap(frame)

    def display_resized_frame2(self, frame):
        self.VisionCamara2.setPixmap(frame)

    def display_resized_frame3(self, frame):
        self.VisionCamara3.setPixmap(frame)
    def display_resized_frame4(self, frame):
        self.VisionCamara4.setPixmap(frame)
    def display_resized_frameB(self, frame):
        self.VisionB.setPixmap(frame)
    def display_resized_frameB_P(self, frame):
        self.VisionB_2.setPixmap(frame)
    
    def display_resized_frame_thermal(self,frame):
        self.VisionB_2.setPixmap(frame)

    def depuracion(self):
        global _FLAG_CONECT

        message=f"{self.textEdit.toPlainText()}"
        if not _FLAG_CONECT:

            # Crear un socket TCP/IP
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)

            # Conectar el socket al puerto donde el servidor está escuchando
            server_address = (IP_ADDR_B, PORT_B)
            print(
                f"Conectando a {server_address[0]} puerto {server_address[1]}")
            s.connect(server_address)
            _FLAG_CONECT = True
        try:
            # Enviar datos
            print(f"Enviando: {message}")
            s.sendall(message.encode('utf-8'))
            # Buscar respuesta
            data = None
            data = s.recv(256)

            if data:

                data = data.decode('utf-8')
                print(f"Recibido: {data}")
                data_n=f'{data}'
                self.label_37.setText(data_n)
                _FLAG_CONECT = False

            else:
                _FLAG_CONECT = False
                self.label_37.setText("Es none")
                print("Es None")
        except IndexError:
            print("No llegó un dato anguloso")

        except:
            print("Cerrando socket")
            s.close()
            _FLAG_CONECT = False
    def stop_recording(self):
        self.label_16.setStyleSheet(f"color: #FF0000")
        self.label_16.setText('DETENIDO')
        global is_taking_pictures

        is_taking_pictures = False

        # self.video_thread1.stop()
        # self.video_thread2.stop()

    def generate_report(self):
        print(f"El texto es {self.empresa_line.text()}")
        fecha_actual = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        fecha = datetime.now().strftime('%d_%m_%Y_%H_%M_%S')

        # Obtener la ruta completa del archivo de imagen
        script_dir = os.path.dirname(os.path.abspath(__file__))
        image_path = os.path.join(script_dir, 'logo.png')

        # Nombre del archivo del documento de Word y PDF
        docx_filename = f'Reportumi_{fecha_actual}.docx'
        pdf_filename = f'Reportumi_{fecha_actual}.pdf'

        print(pdf_filename)

        document = Document()
        document.add_paragraph().alignment = 1
        document.paragraphs[-1].add_run().add_picture(image_path, width=Inches(2.0))
        # Añadir el título centrado
        title = document.add_heading('Reportumi', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Agregar texto con formato
        p = document.add_paragraph()
        p.add_run('Mapa: ').bold = True
        p.add_run(f'{self.comboMapa.currentText()}')

        p = document.add_paragraph()
        p.add_run('Area: ').bold = True
        p.add_run(f'{self.comboArea.currentText()}')

        
        p = document.add_paragraph()
        p.add_run('Fecha de generación de informe: ').bold = True
        p.add_run(f"{fecha_actual}").italic = True

        p = document.add_paragraph()
        p.add_run('Colaborador: ').bold = True
        p.add_run(f'{self.comboCodigo.currentText()}')

        p = document.add_paragraph()
        p.add_run('Turno: ').bold = True
        p.add_run(f'{self.comboTurno.currentText()}')

        p = document.add_paragraph()
        p.add_run('Comentarios: ').bold = True
        p.add_run(f'{self.comentarios_edit.toPlainText()}')

        # Añadir el título del registro de inspección centrado
        inspection_title = document.add_heading('Registro de Inspección', level=1)
        inspection_title.alignment = WD_ALIGN_PARAGRAPH.CENTER


        #document.add_page_break()
        nameDocument=f"Reportumi ({fecha})"

        document.save(f"{nameDocument}.docx")


        # Proporcionar los nombres de archivo para los archivos de Word y PDF
        docx_file = f'{nameDocument}.docx'
        pdf_file = f'{nameDocument}.pdf'

        # Convertir el documento de Word a PDF
        convert(docx_file, pdf_file)
#########################################################################################################################################################################

#########################################################################################################################################################################
# FUNCION PARA AGRANDAR Y ACHICAR LA SENAL DE VIDEO AL AGRANDAR Y ACHICAR LA PANTALLA
    def update_video(self, buffer, id):
        # Load the image from bytes
        pixmap = QPixmap()
        pixmap.loadFromData(buffer)

        # Set the pixmap on the label
        if id == 1:
            self.VisionCamara1.setPixmap(pixmap)
        elif id == 2:
            self.VisionCamara2.setPixmap(pixmap)
        elif id == 3:
            self.VisionCamara3.setPixmap(pixmap)
        elif id == 4:
            self.VisionCamara4.setPixmap(pixmap)

    def maximize_video(self):
        global height
        global width

        height = 540

        width = 960

    def minimize_video(self):
        global height
        global width

        height = 540
        width = 960

    def closeEvent(self, event):
        # Stop the video thread before closing the window
        # self.video_thread1.stop()
        # self.video_thread1.wait()
        # self.video_thread2.stop()
        # self.video_thread2.wait()
        global is_runing
        is_runing = False
        exit_flag.set()
        #self.video_thread4.stop()
        super().closeEvent(event)

    def Autonomous(self):
        global FLAG_AUTO
        if not FLAG_AUTO:
            self.label_mode.setText("Modo Autonomo")
            pixmap = QPixmap("iconos/automode.svg")
            pixmap = pixmap.scaled(60, 60, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
            icon = QIcon(pixmap)
            self.btn_toggle.setIcon(icon)
            self.btn_toggle.setIconSize(QtCore.QSize(60, 60))

            
            

            FLAG_AUTO=True
        else:
            self.label_mode.setText("Modo Teleoperado")
            pixmap = QPixmap("iconos/Toggle.svg")
            pixmap = pixmap.scaled(60, 60, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
            icon = QIcon(pixmap)
            self.btn_toggle.setIcon(icon)
            self.btn_toggle.setIconSize(QtCore.QSize(60, 60))
            FLAG_AUTO=False
        
        pixmap = QPixmap(MAP_PIC)
        pixmap = pixmap.scaled(400, 250, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
        icon = QIcon(pixmap)
        self.btn_mapa.setIcon(icon)
        self.btn_mapa.setIconSize(QtCore.QSize(400, 250))
        pixmap = QPixmap(MAP_PIC)
        pixmap = pixmap.scaled(400, 250, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
        icon = QIcon(pixmap)
        self.btn_mapa_2.setIcon(icon)
        self.btn_mapa_2.setIconSize(QtCore.QSize(400, 250))
        if not FLAG_AUTO:
            self.frame_mapa.setMaximumSize(0, 0)
            self.frame_general.setMaximumSize(10000, 10000)
            self.frame_general_auto.setMaximumSize(0, 0)
        else:
            self.frame_mapa.setMaximumSize(0, 0)
            self.frame_general.setMaximumSize(0, 0)
            self.frame_general_auto.setMaximumSize(10000, 10000)

    def add_items_ComboBox(self,comboBox, csv_file_path, column_name):
        try:
            with open(csv_file_path, newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                for row in reader:
                    if column_name in row:
                        comboBox.addItem(row[column_name])
                    else:
                        print(f"Column '{column_name}' not found in CSV file.")
                        break
        except Exception as e:
            print(f"Error reading CSV file: {e}")
#########################################################################################################################################################################
# FUNCION PARA REAJUSTAR DISTRIBUCION DE FUNCIONES EN LA PANTALLA PRINCIPAL
    def GoGenerarReport(self):
        #Frames Generales 
        self.main_body_left.setMaximumSize(10000, 100000) #Contiene frame_left_up y frame_left_down
        self.main_frames_login.setMaximumSize(0, 0) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(0,0) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(10000,10000) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas
        self.frame_first.setMaximumSize(0,0)
        self.frame_mapa.setMaximumSize(0, 0)
        self.frame_secret.setMaximumSize(0,0)
        self.frame_general.setMaximumSize(0, 0)
        self.frame_arm.setMaximumSize(0, 0)
        self.frame_posicion.setMaximumSize(0,0)
        self.frame_arch.setMaximumSize(0, 0)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(0, 0)
        self.frame_estadisticas.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(10000, 10000)
        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)



    def GoEstadisticas(self):
        #Frames Generales 
        self.main_body_left.setMaximumSize(10000, 100000) #Contiene frame_left_up y frame_left_down
        self.main_frames_login.setMaximumSize(0, 0) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(0,0) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(10000,10000) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas
        self.frame_first.setMaximumSize(0,0)
        self.frame_mapa.setMaximumSize(0, 0)
        self.frame_secret.setMaximumSize(0,0)
        self.frame_general.setMaximumSize(0, 0)
        self.frame_arm.setMaximumSize(0, 0)
        self.frame_posicion.setMaximumSize(0,0)
        self.frame_arch.setMaximumSize(0, 0)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(0, 0)
        self.frame_estadisticas.setMaximumSize(10000, 10000)
        self.frame_generar_report.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(0, 0)
        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)
    # Funcion que distribuye funcion de posicion, control y angulos
    def distribuir_secret(self):

        
        self.movie = QMovie("iconos\marcianito.gif") 
        self.label_gif.setMovie(self.movie)
        self.label_gif.setAlignment(Qt.AlignCenter)
        self.movie.setScaledSize(QSize().scaled(2300, 780, Qt.KeepAspectRatio))
        #self.label_gif.setMaximumSize(100000000, 10000000) 
        self.movie.start() 

        #Frames Generales 
        self.main_body_left.setMaximumSize(10000, 10000) #Contiene frame_left_up
        self.main_frames_login.setMaximumSize(0, 0) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(2000,1240) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(0,0) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas
        self.frame_first.setMaximumSize(0,0)
        self.frame_mapa.setMaximumSize(0, 0)
        self.frame_secret.setMaximumSize(100000,100000)
        self.frame_general.setMaximumSize(0, 0)
        self.frame_arm.setMaximumSize(0, 0)
        self.frame_posicion.setMaximumSize(0,0)
        self.frame_arch.setMaximumSize(0, 0)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(0, 0)
        self.frame_estadisticas.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(0, 0)
        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)
    #################################################################
    def distribuir_arm(self):

        #Frames Generales 
        self.main_body_left.setMaximumSize(10000, 1100) #Contiene frame_left_up
        self.main_frames_login.setMaximumSize(0, 0) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(2000,1240) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(0,0) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas
        self.frame_first.setMaximumSize(0,0)
        self.frame_mapa.setMaximumSize(0, 0)
        self.frame_secret.setMaximumSize(0,0)
        self.frame_general.setMaximumSize(0, 0)
        self.frame_arm.setMaximumSize(10000, 10000)
        self.frame_posicion.setMaximumSize(0,0)
        self.frame_arch.setMaximumSize(0, 0)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(0, 0)
        self.frame_estadisticas.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(0, 0)
        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)
    #################################################################

    def distribuir_general(self):
        global FLAG_MAPA
        #Frames Generales 
        self.main_body_left.setMaximumSize(0, 0) #Contiene frame_left_up
        self.main_frames_login.setMaximumSize(10000, 10000) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(0,0) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(0,0) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas
        self.frame_first.setMaximumSize(0,0)
        self.frame_secret.setMaximumSize(0,0)
        self.frame_arm.setMaximumSize(0, 0)
        self.frame_posicion.setMaximumSize(10000,10040)
        self.frame_arch.setMaximumSize(0, 0)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(0, 0)
        self.frame_estadisticas.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(0, 0)
        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)
        if not FLAG_MAPA:
            self.frame_mapa.setMaximumSize(100000, 100000)
            self.frame_general.setMaximumSize(0, 0)
            self.frame_general_auto.setMaximumSize(0, 0)
        else:
            
            # self.btn_mapa2.setIcon(icon)
            # self.btn_mapa2.setIconSize(QtCore.QSize(400, 250))
            if not FLAG_AUTO:
                self.frame_mapa.setMaximumSize(0, 0)
                self.frame_general.setMaximumSize(10000, 10000)
                self.frame_general_auto.setMaximumSize(0, 0)
            else:
                self.frame_mapa.setMaximumSize(0, 0)
                self.frame_general.setMaximumSize(0, 0)
                self.frame_general_auto.setMaximumSize(10000, 10000)
                # self.btn_VisionCamara1_3.lower()
                # self.setLayout(self.frame_cam_map)
                # self.frame_cam_map().removeWidget(self.btn_VisionCamara1_3)
                # self.frame_cam_map().removeWidget(self.btn_map)
                # self.frame_cam_map().addWidget(self.btn_map)
                # self.frame_cam_map().addWidget(self.btn_VisionCamara1_3)

       
    #################################################################

    def distribuir_document(self):

        #Frames Generales 
        self.main_body_left.setMaximumSize(10000, 1100) #Contiene frame_left_up
        self.main_frames_login.setMaximumSize(0, 0) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(0,0) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(10000,10000) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas
        self.frame_first.setMaximumSize(0,0)
        self.frame_mapa.setMaximumSize(0, 0)
        self.frame_secret.setMaximumSize(0,0)
        self.frame_general.setMaximumSize(0, 0)
        self.frame_arm.setMaximumSize(0, 0)
        self.frame_posicion.setMaximumSize(0,0)
        self.frame_arch.setMaximumSize(0, 0)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(100000, 100000)
        self.frame_estadisticas.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(0, 0)
        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)
    #################################################################
    # Funcion que distribuye funcion de camara

    def distribuir_posicion(self):

        #Frames Generales 
        self.main_body_left.setMaximumSize(10000, 1100) #Contiene frame_left_up
        self.main_frames_login.setMaximumSize(0, 0) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(2000,1240) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(0,0) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas
        self.frame_first.setMaximumSize(0,0)
        self.frame_mapa.setMaximumSize(0, 0)
        self.frame_secret.setMaximumSize(0,0)
        self.frame_general.setMaximumSize(0, 0)
        self.frame_arm.setMaximumSize(0, 0)
        self.frame_posicion.setMaximumSize(10000,10040)
        self.frame_arch.setMaximumSize(0, 0)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(0, 0)
        self.frame_estadisticas.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(0, 0)
        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)
###############################################

    def GoGeneral(self,name,area_map, img):
        global FLAG_MAPA, MAP_NAME,MAP_AREA,MAP_PIC
        FLAG_MAPA=True
        MAP_NAME=name
        MAP_AREA=area_map
        MAP_PIC=img

        pixmap = QPixmap(MAP_PIC)
        pixmap = pixmap.scaled(400, 250, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
        icon = QIcon(pixmap)
        self.btn_mapa.setIcon(icon)
        self.btn_mapa.setIconSize(QtCore.QSize(400, 250))
        pixmap = QPixmap(MAP_PIC)
        pixmap = pixmap.scaled(400, 250, QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation)
        icon = QIcon(pixmap)
        self.btn_mapa_2.setIcon(icon)
        self.btn_mapa_2.setIconSize(QtCore.QSize(400, 250))

        #Frames Generales 
        self.main_body_left.setMaximumSize(0, 0) #Contiene frame_left_up
        self.main_frames_login.setMaximumSize(10000, 10000) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(0,0) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(0,0) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas
        self.frame_first.setMaximumSize(0,0)
        self.frame_mapa.setMaximumSize(0, 0)
        self.frame_general.setMaximumSize(10000, 10000)
        self.frame_secret.setMaximumSize(0,0)
        self.frame_arm.setMaximumSize(0, 0)
        self.frame_posicion.setMaximumSize(10000,10040)
        self.frame_arch.setMaximumSize(0, 0)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(0, 0)
        self.frame_estadisticas.setMaximumSize(0, 0)

        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)
    # Funcion que distribuye funcion de archivos

    def distribuir_archivos(self):

        #Frames Generales 
        self.main_body_left.setMaximumSize(10000, 1100) #Contiene frame_left_up
        self.main_frames_login.setMaximumSize(0, 0) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(0,0) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(100000,10000) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas
        self.frame_first.setMaximumSize(0,0)
        self.frame_mapa.setMaximumSize(0, 0)
        self.frame_secret.setMaximumSize(0,0)
        self.frame_general.setMaximumSize(0, 0)
        self.frame_arm.setMaximumSize(0, 0)
        self.frame_posicion.setMaximumSize(0,0)
        self.frame_arch.setMaximumSize(10000, 10000)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(0, 0)
        self.frame_estadisticas.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(0, 0)

        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)

        self.load_images_galery()
    # Funcion para reiniciar distribucion de espacios

    def reset_button(self):
        #Frames Generales 
        self.main_body_left.setMaximumSize(0, 0) #Contiene frame_left_up
        self.main_frames_login.setMaximumSize(100000, 10000) #Contiene a frame_general y a frame_mapa
        self.frame_left_up.setMaximumSize(0,0) #Contiene a frame_arm, frame_position y frame_secret
        self.frame_left_down.setMaximumSize(0,0) #Contiene a frame_arch, frame_documents y derivados del reporte

        #Frames Ventanas

        self.frame_mapa.setMaximumSize(0, 0)
        self.frame_first.setMaximumSize(10000,10000)
        self.frame_secret.setMaximumSize(0,0)
        self.frame_general.setMaximumSize(0, 0)
        self.frame_arm.setMaximumSize(0, 0)
        self.frame_posicion.setMaximumSize(0,0)
        self.frame_arch.setMaximumSize(0, 0)

        #Frames Reportes
        self.frame_menu_report.setMaximumSize(0, 0)
        self.frame_estadisticas.setMaximumSize(0, 0)
        self.frame_generar_report.setMaximumSize(0, 0)
        #Labels Camaras
        self.VisionCamara1.setMinimumSize(960,570)
        self.VisionCamara1.setMaximumSize(960,570)
        self.VisionCamara2.setMinimumSize(250, 150)
        self.VisionCamara2.setMaximumSize(250, 150)

        



    def saveimages(self):
        global is_saving
                
        root = Tk()
        root.withdraw()  # Ocultar la ventana principal de Tkinter
        folder_path = filedialog.askdirectory()
        output_folder = folder_path

        is_saving = True
        
        
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        # Captura el video
        cap = cv2.VideoCapture(0)

        frame_count = 0
        while is_saving:
            ret, frame = cap.read()
            if not ret:
                break
            # Guarda cada frame como una imagen
            frame_filename = os.path.join(output_folder, f"frame_{frame_count:04d}.png")
            cv2.imwrite(frame_filename, frame)
            frame_count += 1

        # Libera la captura de video
        cap.release()
        print(f"Se han guardado {frame_count} frames en la carpeta {output_folder}")

    def stop_saving(self):
        global is_saving
        is_saving = False
#########################################################################################################################################################################
# FUNCION MENU DESLIZANTE

    def slideLeftMenu(self):
        # Obtenemos la medida del ancho del menu
        width = self.slide_menu_container.width()

        # Si esta minimized
        if width == 0:
            # Expand menu
            newWidth = 300
            self.open_close_side_bar_btn.setIcon(
                QtGui.QIcon(u"Iconos/left.png"))

        # Si el menu esta abierto
        else:
            # Restore menu
            newWidth = 0
            self.open_close_side_bar_btn.setIcon(
                QtGui.QIcon(u"Iconos/menu.png"))

        # Transision de animacion
        self.animation = QPropertyAnimation(
            self.slide_menu_container, b"maximumWidth")  # Animate minimumWidht
        self.animation.setDuration(450)
        # Start value is the current menu width
        self.animation.setStartValue(width)
        self.animation.setEndValue(newWidth)  # end value is the new menu width
        self.animation.setEasingCurve(QtCore.QEasingCurve.InOutQuart)
        self.animation.start()


#########################################################################################################################################################################
# FUNCION MOVER VENTANA

    # Evento de mouse

    def mousePressEvent(self, event):
        # Get the current position of the mouse
        self.clickPosition = event.globalPos()
        # We will use this value to move the window

